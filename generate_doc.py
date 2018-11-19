from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches

#打开文档
document = Document()

#加入不同等级的标题
document.add_heading('Document Title',0)
document.add_heading('二级标题',1)
document.add_heading('二级标题',2)

#添加文本
paragraph = document.add_paragraph('添加了文本')
#设置字号
run = paragraph.add_run('设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'

#设置中文字体
run = paragraph.add_run('设置中文字体，')
run.font.name='宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

#设置斜体
run = paragraph.add_run('斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run('粗体').bold = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    '有序列表元素1',style='List Number'
)
document.add_paragraph(
    '有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    '无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    '无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
document.add_picture('jdb.jpg',width=Inches(1.25))

#增加表格
table = document.add_table(rows=3,cols=3)
hdr_cells=table.rows[0].cells
hdr_cells[0].text="第一列"
hdr_cells[1].text="第二列"
hdr_cells[2].text="第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = 'aerszvfdgx'
hdr_cells[2].text = 'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '3'
hdr_cells[1].text = 'cafdwvaef'
hdr_cells[2].text = 'aabs zfgf'

#增加分页
document.add_page_break()

#保存文件
document.save('demo1.docx')
